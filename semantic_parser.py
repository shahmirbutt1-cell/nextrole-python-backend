"""
semantic_parser.py — format-agnostic resume parser

Detection improvements over the previous version:
  - Font size used as a header signal (not just bold/all-caps)
  - Underline used as a header signal
  - Header word-count limit raised from 6 → 10
  - Paragraph indentation used as a bullet signal
  - Expanded bullet character set (›, ◦, ▪, ○, →, ✓, etc.)
  - Numbered / lettered bullet detection (1. / a) / i.)
  - Custom List-variant style names recognised
  - Greatly expanded _HEADER_RULES (Internships, Projects, Volunteer, etc.)
  - _is_role_header accepts em-dash, no-space dashes, and "Present" as signals
  - Font-size-based role header detection (larger text = title)
  - Two-column table layouts read left-column-first, right-column-second
  - Textbox content extracted where possible
  - Entire document sent to GPT as a last-resort fallback
  - All items in resume_model are JSON-serializable (no raw docx objects)
  - Raw docx objects stored separately in _paragraph_refs keyed by ref_id
"""

import json
import logging
import re
from statistics import median
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.oxml.ns import qn

# ─────────────────────────────────────────────────────────────────────────────
# Module-level constants — built once, never rebuilt per-call
# ─────────────────────────────────────────────────────────────────────────────

# Maps section category → list of regex patterns matched against normalised text
_HEADER_RULES: Dict[str, List[str]] = {
    "summary": [
        r"^summary$",
        r"^professional summary$",
        r"^executive summary$",
        r"^profile$",
        r"^professional profile$",
        r"^career summary$",
        r"^career profile$",
        r"^about$",
        r"^about me$",
        r"^objective$",
        r"^career objective$",
        r"^personal statement$",
        r"^overview$",
        r"^professional overview$",
        r"^introduction$",
    ],
    "skills": [
        r"^skills$",
        r"^skill set$",
        r"^technical skills$",
        r"^technology skills$",
        r"^core competencies$",
        r"^competencies$",
        r"^key competencies$",
        r"^tools$",
        r"^tools and technologies$",
        r"^technologies$",
        r"^technical expertise$",
        r"^areas of expertise$",
        r"^expertise$",
        r"^qualifications$",
        r"^key qualifications$",
        r"^key skills$",
        r"^strengths$",
        r"^core skills$",
        r"^professional skills$",
        r"^languages$",
        r"^programming languages$",
        r"^software$",
        r"^software skills$",
        r"^certifications$",
        r"^certificates$",
        r"^licenses and certifications$",
        r"^licenses$",
        r"^technical proficiencies$",
        r"^proficiencies$",
    ],
    "experience": [
        r"^experience$",
        r"^work experience$",
        r"^professional experience$",
        r"^relevant experience$",
        r"^related experience$",
        r"^industry experience$",
        r"^employment$",
        r"^employment history$",
        r"^work history$",
        r"^career history$",
        r"^career experience$",
        r"^job history$",
        r"^positions held$",
        r"^professional background$",
        r"^internships?$",
        r"^internship experience$",
        r"^volunteer(ing)?$",
        r"^volunteer experience$",
        r"^volunteer work$",
        r"^community involvement$",
        r"^leadership experience$",
        r"^research experience$",
        r"^teaching experience$",
        r"^consulting experience$",
        r"^freelance experience$",
        r"^projects?$",
        r"^key projects?$",
        r"^relevant projects?$",
        r"^selected projects?$",
        r"^notable projects?$",
    ],
    "education": [
        r"^education$",
        r"^education and training$",
        r"^academic background$",
        r"^academic history$",
        r"^academic qualifications$",
        r"^educational background$",
        r"^educational history$",
        r"^degrees?$",
        r"^training$",
        r"^training and development$",
        r"^professional development$",
        r"^courses?$",
        r"^coursework$",
        r"^continuing education$",
        r"^awards?$",
        r"^awards and honors?$",
        r"^honors?$",
        r"^scholarships?$",
        r"^publications?$",
        r"^presentations?$",
        r"^conferences?$",
    ],
}

# Bullet characters beyond the basic set
_BULLET_CHARS = frozenset("•-–—*›◦▪▸▹○●◆◇→✓✔►")

# Regex for numbered / lettered bullets: "1.", "2)", "a.", "b)", "i.", "ii."
_NUMBERED_BULLET_RE = re.compile(
    r"^(\d{1,2}[.)]\s|[a-z][.)]\s|[ivxlcdm]{1,4}[.)]\s)", re.IGNORECASE
)


# ─────────────────────────────────────────────────────────────────────────────
# Helper: extract font size from a paragraph (in half-points, or None)
# ─────────────────────────────────────────────────────────────────────────────

def _paragraph_font_size(paragraph) -> Optional[int]:
    """
    Return the dominant font size (half-points) in a paragraph, or None.
    Checks run-level size first, then paragraph style, then document default.
    """
    sizes = []
    for run in paragraph.runs:
        sz = run.font.size
        if sz:
            sizes.append(int(sz))
    if sizes:
        return max(sizes)

    # Fall back to paragraph style
    try:
        style_sz = paragraph.style.font.size
        if style_sz:
            return int(style_sz)
    except Exception:
        pass

    return None


def _paragraph_is_underlined(paragraph) -> bool:
    """Return True if any run in the paragraph is underlined."""
    return any(run.underline for run in paragraph.runs)


def _paragraph_indent_twips(paragraph) -> int:
    """Return the left indent in twips (0 if none)."""
    try:
        ind = paragraph.paragraph_format.left_indent
        if ind is not None:
            return int(ind)
    except Exception:
        pass
    return 0


# ─────────────────────────────────────────────────────────────────────────────
# Helper: extract text from textboxes embedded in the document XML
# ─────────────────────────────────────────────────────────────────────────────

def _extract_textbox_paragraphs(doc) -> List[str]:
    """
    python-docx doesn't expose textboxes through its API.
    We reach into the raw XML to pull text from <w:txbxContent> elements.
    """
    texts = []
    try:
        body = doc.element.body
        for txbx in body.iter(qn("w:txbxContent")):
            for p in txbx.iter(qn("w:p")):
                parts = []
                for t in p.iter(qn("w:t")):
                    if t.text:
                        parts.append(t.text)
                text = "".join(parts).strip()
                if text:
                    texts.append(text)
    except Exception as e:
        logging.warning(f"Textbox extraction failed: {e}")
    return texts


# ─────────────────────────────────────────────────────────────────────────────
# Main parser class
# ─────────────────────────────────────────────────────────────────────────────

class SemanticResumeParser:
    """
    Parse a .docx resume into structured, JSON-serializable data.

    Accepts either a file path string or a pre-loaded python-docx Document.
    Uses GPT for section header classification when an OpenAI client is
    provided, falling back to regex rules otherwise.

    Raw docx paragraph objects are stored in self._paragraph_refs (keyed by
    ref_id = id(paragraph)) so that ResumeTailorEngine can write back changes
    without the resume_model needing to hold non-serializable objects.

    Usage:
        parser = SemanticResumeParser("resume.docx", openai_client=client)
        model  = parser.parse()
        # model keys: summary, skills, experience, education
        # raw objects: parser._paragraph_refs[ref_id]
    """

    def __init__(self, source: Union[str, Any], openai_client=None):
        self.client = openai_client
        self.paragraphs: List[Dict[str, Any]] = []
        self.sections: Dict[str, List[Dict[str, Any]]] = {}
        self._paragraph_refs: Dict[int, Any] = {}

        self.resume_model: Dict[str, Any] = {
            "summary":    None,
            "skills":     [],
            "experience": [],
            "education":  [],
        }

        if hasattr(source, "paragraphs") and hasattr(source, "tables"):
            self.file_path = None
            self.doc = source
        else:
            self.file_path = source
            self.doc = Document(source)

        # Computed during _extract_paragraphs; used for relative size checks
        self._median_font_size: Optional[int] = None

    # ──────────────────────────────────────────────────────────────────────────
    # PUBLIC ENTRY
    # ──────────────────────────────────────────────────────────────────────────

    def parse(self) -> Dict[str, Any]:
        """
        Parse the loaded resume into structured data.

        Returns a JSON-serializable dict with keys:
            summary, skills, experience, education

        Raw docx paragraph objects are in self._paragraph_refs.
        """
        if self.resume_model.get("_parsed"):
            return self.resume_model

        self._extract_paragraphs()
        self._detect_sections()
        self._parse_summary()
        self._parse_skills()
        self._parse_experience()
        self._parse_education()

        self.resume_model["_parsed"] = True
        return self.resume_model

    # ──────────────────────────────────────────────────────────────────────────
    # STEP 1 — EXTRACT PARAGRAPHS
    # ──────────────────────────────────────────────────────────────────────────

    def _extract_paragraphs(self):
        """
        Build self.paragraphs in document reading order.

        Key improvements:
        - Table cell paragraphs are de-duplicated (python-docx includes them
          in both doc.paragraphs and doc.tables).
        - Font size and indent are captured per paragraph for richer detection.
        - Median font size is computed so we can flag larger-than-normal text
          as a header signal.
        - Textbox text is appended at the end (read-only; no ref_id since
          textbox paragraphs can't be written back via python-docx).
        """
        def capture(p) -> Dict[str, Any]:
            ref_id = id(p)
            self._paragraph_refs[ref_id] = p
            return {
                "text":      p.text.strip(),
                "style":     p.style.name if p.style else "",
                "bold":      any(run.bold for run in p.runs),
                "underline": _paragraph_is_underlined(p),
                "font_size": _paragraph_font_size(p),
                "indent":    _paragraph_indent_twips(p),
                "ref_id":    ref_id,
            }

        # Collect table-cell paragraph objects so we can skip them in the
        # main doc.paragraphs loop (they appear in both places)
        table_paragraph_objects = {
            p
            for table in self.doc.tables
            for row in table.rows
            for cell in row.cells
            for p in cell.paragraphs
        }

        self.paragraphs = []

        # Main body (non-table) paragraphs
        for p in self.doc.paragraphs:
            if p.text.strip() and p not in table_paragraph_objects:
                self.paragraphs.append(capture(p))

        # Table cell paragraphs — left-to-right, top-to-bottom
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            self.paragraphs.append(capture(p))

        # Compute median font size across all paragraphs that have one
        sizes = [p["font_size"] for p in self.paragraphs if p["font_size"]]
        self._median_font_size = int(median(sizes)) if sizes else None

        # Textbox content (no ref_id — can't write back)
        for text in _extract_textbox_paragraphs(self.doc):
            self.paragraphs.append({
                "text":      text,
                "style":     "",
                "bold":      False,
                "underline": False,
                "font_size": None,
                "indent":    0,
                "ref_id":    None,   # textbox — not writable
            })

    # ──────────────────────────────────────────────────────────────────────────
    # STEP 2 — SECTION DETECTION
    # ──────────────────────────────────────────────────────────────────────────

    def _is_header_candidate(self, p: Dict[str, Any]) -> bool:
        """
        Return True if a paragraph looks like a section header.

        Signals (any one is sufficient):
          - Bold text with ≤ 10 words
          - ALL CAPS text with ≤ 10 words
          - Underlined text with ≤ 10 words
          - Font size noticeably larger than the document median
          - Text matches a known header pattern directly (belt-and-suspenders)
        """
        text  = p["text"]
        words = text.split()

        if not text or len(words) > 10:
            return False

        if p["bold"] or text.isupper() or p["underline"]:
            return True

        # Larger-than-normal font (≥ 20% above median)
        if self._median_font_size and p["font_size"]:
            if p["font_size"] >= self._median_font_size * 1.2:
                return True

        # Direct match against known patterns (catches plain-text resumes)
        normalized = re.sub(r"[^a-z\s]", "", text.lower()).strip()
        for patterns in _HEADER_RULES.values():
            if any(re.match(pat, normalized) for pat in patterns):
                return True

        return False

    def _identify_header_candidates(self) -> List[Dict[str, Any]]:
        return [p for p in self.paragraphs if self._is_header_candidate(p)]

    def _validate_headers_with_gpt(
        self, candidates: List[Dict[str, Any]]
    ) -> Optional[Dict[str, str]]:
        """
        Ask GPT to classify each candidate header.
        Returns {header_text: category} or None on any failure.
        """
        if not self.client or not candidates:
            return None

        headers_joined = "\n".join(p["text"] for p in candidates)

        prompt = f"""You are a resume structure classification engine.

Classify each line below into ONE of these categories:
- summary
- skills
- experience
- education
- none

Return ONLY valid JSON: {{"Header Text": "category"}}

Headers:
{headers_joined}"""

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You classify resume section headers. Return only valid JSON."},
                    {"role": "user",   "content": prompt},
                ],
                timeout=10,
            )
        except Exception as e:
            logging.warning(f"GPT header classification request failed: {e}")
            return None

        content = response.choices[0].message.content.strip()
        # Strip accidental markdown code fences
        content = re.sub(r"^```[a-z]*\n?", "", content).rstrip("` \n")

        try:
            parsed = json.loads(content)
        except json.JSONDecodeError as e:
            logging.warning(
                f"GPT returned invalid JSON for header classification: {e}. "
                f"Raw (first 300 chars): {content[:300]}"
            )
            return None

        if not isinstance(parsed, dict):
            logging.warning(f"GPT classification was not a dict: {type(parsed)}")
            return None

        allowed = {"summary", "skills", "experience", "education", "none"}
        bad = [v for v in parsed.values() if v not in allowed]
        if bad:
            logging.warning(f"GPT returned unexpected labels: {bad}. Falling back to regex.")
            return None

        return parsed

    def _classify_header_fallback(self, text: str) -> Optional[str]:
        """Classify a header using regex rules. Returns category or None."""
        normalized = re.sub(r"[^a-z\s]", "", text.lower()).strip()
        for label, patterns in _HEADER_RULES.items():
            if any(re.match(pat, normalized) for pat in patterns):
                return label
        return None

    def _gpt_classify_full_document(self) -> Optional[Dict[str, str]]:
        """
        Last-resort: send the entire document text to GPT and ask it to
        identify section boundaries. Used when candidate detection finds
        nothing (e.g. plain-text or visually-styled resumes).
        """
        if not self.client:
            return None

        all_text = "\n".join(p["text"] for p in self.paragraphs[:80])

        prompt = f"""Below is the text of a resume. Identify which lines are section headers
and classify each into: summary, skills, experience, education, or none.

Return ONLY valid JSON: {{"exact line text": "category"}}

Resume text:
{all_text}"""

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You identify resume section headers. Return only valid JSON."},
                    {"role": "user",   "content": prompt},
                ],
                timeout=15,
            )
        except Exception as e:
            logging.warning(f"GPT full-document classification failed: {e}")
            return None

        content = response.choices[0].message.content.strip()
        content = re.sub(r"^```[a-z]*\n?", "", content).rstrip("` \n")

        try:
            parsed = json.loads(content)
        except json.JSONDecodeError as e:
            logging.warning(f"GPT full-document classification returned invalid JSON: {e}")
            return None

        allowed = {"summary", "skills", "experience", "education", "none"}
        return {k: v for k, v in parsed.items() if v in allowed} or None

    def _detect_sections(self):
        section_map: Dict[str, List[Dict]] = {
            "summary":    [],
            "skills":     [],
            "experience": [],
            "education":  [],
        }

        candidates = self._identify_header_candidates()

        # Try GPT on candidates first
        classification = self._validate_headers_with_gpt(candidates)

        # Fall back to regex on candidates
        if not classification:
            classification = {
                p["text"]: self._classify_header_fallback(p["text"]) or "none"
                for p in candidates
            }

        # If we detected zero real sections, escalate to full-document GPT
        found_sections = {v for v in classification.values() if v != "none"}
        if not found_sections and self.client:
            logging.info("No sections found via candidates — trying full-document GPT classification")
            full_doc_result = self._gpt_classify_full_document()
            if full_doc_result:
                classification = full_doc_result

        current_section = None
        for p in self.paragraphs:
            text = p["text"]
            if text in classification:
                label = classification[text]
                current_section = label if label in section_map else None
                continue
            if current_section:
                section_map[current_section].append(p)

        self.sections = section_map

        # Log what we found for debugging
        for sec, items in section_map.items():
            logging.info(f"Section '{sec}': {len(items)} paragraphs")

    # ──────────────────────────────────────────────────────────────────────────
    # STEP 3 — SUMMARY
    # ──────────────────────────────────────────────────────────────────────────

    def _parse_summary(self):
        paras = self.sections.get("summary", [])
        if not paras:
            return

        combined = " ".join(p["text"] for p in paras)
        self.resume_model["summary"] = {
            "text": combined,
            "paragraphs": [
                {"text": p["text"], "ref_id": p["ref_id"]}
                for p in paras
            ],
        }

    # ──────────────────────────────────────────────────────────────────────────
    # STEP 4 — SKILLS
    # ──────────────────────────────────────────────────────────────────────────

    def _parse_skills(self):
        for p in self.sections.get("skills", []):
            self.resume_model["skills"].append({
                "text":   p["text"],
                "ref_id": p["ref_id"],
            })

    # ──────────────────────────────────────────────────────────────────────────
    # STEP 5 — EXPERIENCE
    # ──────────────────────────────────────────────────────────────────────────

    def _parse_experience(self):
        paras  = self.sections.get("experience", [])
        roles  = []
        current_role = None

        for p in paras:
            if self._is_role_header(p):
                if current_role:
                    roles.append(current_role)
                current_role = {
                    "title":        p["text"],
                    "title_ref_id": p["ref_id"],
                    "bullets":      [],
                }
                continue

            if current_role and self._is_bullet(p):
                current_role["bullets"].append({
                    "text":   p["text"],
                    "ref_id": p["ref_id"],
                })

        if current_role:
            roles.append(current_role)

        self.resume_model["experience"] = roles

    def _is_role_header(self, p: Dict[str, Any]) -> bool:
        """
        Determine whether a paragraph is a job role / position header.

        Improvements over previous version:
        - Checks _is_bullet first (bold bullets are not headers)
        - Accepts em-dashes and no-space dashes
        - Accepts "Present" as a date signal (no year needed)
        - Accepts font-size-based detection (larger = title)
        - Accepts standalone bold/underline short lines as role headers
          when inside the experience section
        """
        if self._is_bullet(p):
            return False

        text = p["text"]

        # Year pattern (strongest signal)
        if re.search(r"\b(19|20)\d{2}\b", text):
            return True

        # "Present" without a year (e.g. "Jan 2022 – Present")
        if re.search(r"\bpresent\b", text, re.IGNORECASE):
            return True

        # Any dash variant as separator — with or without spaces
        if re.search(r"\s[-–—]\s|[-–—]", text) and (p["bold"] or p["underline"]):
            return True

        # Font-size signal: noticeably larger than median
        if self._median_font_size and p["font_size"]:
            if p["font_size"] >= self._median_font_size * 1.15 and len(text.split()) <= 12:
                return True

        # Bold/underline short line (≤ 10 words) — likely a company or title
        if (p["bold"] or p["underline"]) and 1 <= len(text.split()) <= 10:
            return True

        return False

    def _is_bullet(self, p: Dict[str, Any]) -> bool:
        """
        Detect bullet points across all common resume formats.

        Signals:
        - Starts with a known bullet character
        - Starts with a numbered/lettered list marker (1. / a) / i.)
        - Style name contains 'list' or 'bullet' (case-insensitive)
        - Non-zero left indent (paragraph is indented — likely a bullet)
        """
        text  = p["text"]
        style = p["style"].lower()

        # Bullet character at start (strip leading whitespace)
        first_char = text.lstrip()[:1]
        if first_char in _BULLET_CHARS:
            return True

        # Numbered / lettered list
        if _NUMBERED_BULLET_RE.match(text.lstrip()):
            return True

        # Style name signals
        if "list" in style or "bullet" in style:
            return True

        # Indented paragraph (threshold: 180 twips ≈ 0.125 inch)
        if p["indent"] >= 180:
            return True

        return False

    # ──────────────────────────────────────────────────────────────────────────
    # STEP 6 — EDUCATION
    # ──────────────────────────────────────────────────────────────────────────

    def _parse_education(self):
        for p in self.sections.get("education", []):
            self.resume_model["education"].append({
                "text":   p["text"],
                "ref_id": p["ref_id"],
            })
