import logging
import os
import re
import uuid

from docx import Document
from fastapi import FastAPI, File, Form, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from openai import OpenAI
from pydantic import BaseModel

from resume_tailor import ResumeTailorEngine
from semantic_parser import SemanticResumeParser

logging.basicConfig(level=logging.INFO)

# ── App must be created BEFORE add_middleware ────────────────────────────────
app = FastAPI(title="NextRole API", version="1.0.0")

# ── CORS ─────────────────────────────────────────────────────────────────────
# Set NEXT_PUBLIC_FRONTEND_URL in Render environment variables to your
# Vercel deployment URL, e.g. https://nextrole.vercel.app
# Falls back to localhost for local development.
_frontend_url = os.getenv("NEXT_PUBLIC_FRONTEND_URL", "http://localhost:3000")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[_frontend_url, "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── OpenAI client ─────────────────────────────────────────────────────────────
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# ── Startup check ─────────────────────────────────────────────────────────────
@app.on_event("startup")
async def startup_check():
    """Fail fast if required environment variables are missing."""
    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError(
            "OPENAI_API_KEY environment variable is not set. "
            "Set it before starting the server."
        )


# ── Health check ──────────────────────────────────────────────────────────────
@app.get("/")
def root():
    return {"status": "NextRole Python backend running"}


# =========================================
# HELPERS
# =========================================

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a .docx file as a single newline-joined string.
    Uses a set of table paragraph objects to avoid double-counting content.
    """
    doc = Document(file_path)

    table_paragraph_objects = {
        p
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for p in cell.paragraphs
    }

    full_text = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() and paragraph not in table_paragraph_objects:
            full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    return "\n".join(full_text)


def extract_keywords(text: str) -> set:
    words = re.findall(r"\b\w+\b", text.lower())
    stopwords = {
        "the", "and", "with", "for", "from", "that", "this", "have", "has",
        "are", "was", "were", "will", "shall", "your", "their", "about",
        "into", "within", "across", "using", "use", "used"
    }
    return set(w for w in words if len(w) > 3 and w not in stopwords)


def role_relevance_score(role: dict, job_keywords: set) -> int:
    role_text = " ".join([b["text"] for b in role["bullets"]]).lower()
    role_words = set(re.findall(r"\b\w+\b", role_text))
    overlap = role_words.intersection(job_keywords)
    if not role_words:
        return 0
    return len(overlap)


def detect_industry_keywords(text: str) -> dict:
    finance_terms = {"finance", "financial", "accounting", "budget", "forecast", "audit"}
    medical_terms = {"medical", "device", "orthopedic", "trauma", "surgical", "hospital"}
    tech_terms    = {"software", "saas", "cloud", "api", "engineering", "data", "ai"}

    text_words = set(re.findall(r"\b\w+\b", text.lower()))

    return {
        "finance": len(text_words.intersection(finance_terms)),
        "medical": len(text_words.intersection(medical_terms)),
        "tech":    len(text_words.intersection(tech_terms)),
    }


# =========================================
# ENDPOINTS
# =========================================

@app.post("/analyze-resume")
async def analyze_resume(file: UploadFile = File(...)):
    """Extract and return the text content of an uploaded .docx resume."""
    import shutil
    tmp_path = f"analyze_{uuid.uuid4().hex}.docx"

    try:
        with open(tmp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        extracted_text = extract_text_from_docx(tmp_path)

        return {
            "file_name":   file.filename,
            "text_length": len(extracted_text),
            "preview":     extracted_text[:300],
            "full_text":   extracted_text,
        }
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


class MatchJobRequest(BaseModel):
    resume_text:     str
    job_description: str


@app.post("/match-job")
async def match_job(data: MatchJobRequest):
    """Score how well a resume matches a job description by keyword overlap."""
    resume_text     = data.resume_text.lower()
    job_description = data.job_description.lower()

    resume_words = set(re.findall(r"\w+", resume_text))
    job_words    = set(re.findall(r"\w+", job_description))

    matching = resume_words.intersection(job_words)
    missing  = job_words.difference(resume_words)

    score = 0 if not job_words else int(len(matching) / len(job_words) * 100)

    return {
        "match_score":     score,
        "matching_skills": list(matching)[:10],
        "missing_skills":  list(missing)[:10],
    }


@app.post("/debug-parse")
async def debug_parse(file: UploadFile = File(...)):
    """Parse an uploaded resume and return the structured model (dev use)."""
    import shutil
    tmp_path = f"debug_{uuid.uuid4().hex}.docx"

    try:
        with open(tmp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        doc    = Document(tmp_path)
        parser = SemanticResumeParser(doc, openai_client=client)
        model  = parser.parse()

        logging.info("==== PARSE DEBUG ====")
        logging.info(f"Summary: {model.get('summary')}")
        logging.info(f"Skills count: {len(model.get('skills', []))}")
        logging.info(f"Experience roles: {len(model.get('experience', []))}")
        logging.info("==== PARSE DEBUG END ====")

        return model
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


@app.post("/tailor-resume-docx-preserve")
async def tailor_resume_docx_preserve(
    file:            UploadFile = File(...),
    job_description: str        = Form(...),
    mode:            str        = Form("balanced"),
):
    """
    Accept a .docx resume and job description, rewrite to match the role,
    and return the tailored .docx. Original formatting is preserved.
    Temp files are always cleaned up after the request.
    """
    import shutil
    from starlette.background import BackgroundTask

    request_id      = uuid.uuid4().hex
    input_filename  = f"input_{request_id}.docx"
    output_filename = f"tailored_{request_id}.docx"

    try:
        with open(input_filename, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        doc          = Document(input_filename)
        parser       = SemanticResumeParser(doc, openai_client=client)
        resume_model = parser.parse()

        logging.info("==== PARSE DEBUG ====")
        logging.info(f"Summary: {resume_model.get('summary')}")
        logging.info(f"Skills count: {len(resume_model.get('skills', []))}")
        logging.info(f"Experience roles: {len(resume_model.get('experience', []))}")
        logging.info("==== PARSE DEBUG END ====")

        tailor_engine = ResumeTailorEngine(
            resume_model    = resume_model,
            paragraph_refs  = parser._paragraph_refs,
            job_description = job_description,
            openai_client   = client,
            mode            = mode,
        )
        tailor_engine.tailor()

        doc.save(output_filename)

        def cleanup():
            for f in [input_filename, output_filename]:
                if os.path.exists(f):
                    os.remove(f)

        return FileResponse(
            output_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="Tailored_Resume.docx",
            background=BackgroundTask(cleanup),
        )

    except Exception:
        for f in [input_filename, output_filename]:
            if os.path.exists(f):
                os.remove(f)
        raise
